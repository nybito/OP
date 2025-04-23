import tkinter as tk
from tkinter import messagebox
from docx import Document
import openpyxl
import math
from module import *

def calculate():
    shape = geometry_type.get()
    try:
        if shape == "Параллелепипед":
            a = float(entry_a.get())
            b = float(entry_b.get())
            h = float(entry_h.get())
            volume = Para(a,b,h)[0]
            surface_area = Para(a,b,h)[1]

        elif shape == "Тетраэдр":
            a = float(entry_a.get())
            volume = Tetra(a)[0]
            surface_area = Tetra(a)[1]

        elif shape == "Шар":
            r = float(entry_r.get())
            volume = Shar(r)[0]
            surface_area = Shar(r)[1]

        else:
            raise ValueError("Неверная фигура")

        material_density = float(entry_density.get())
        mass = volume * material_density

        result_text = f"Объем: {volume:.2f}\nПлощадь поверхности: {surface_area:.2f}\nМасса: {mass:.2f}"
        result_label.config(text=result_text)

    except ValueError as ve:
        messagebox.showerror("Ошибка", str(ve))
def save_report():
    shape = geometry_type.get()
    volume = result_label.cget("text").split("\n")[0].split(":")[1].strip()
    surface_area = result_label.cget("text").split("\n")[1].split(":")[1].strip()
    mass = result_label.cget("text").split("\n")[2].split(":")[1].strip()

    doc = Document()
    doc.add_heading(f'Отчет по {shape}', level=1)
    doc.add_paragraph(f'Объем: {volume}')
    doc.add_paragraph(f'Площадь поверхности: {surface_area}')
    doc.add_paragraph(f'Mасса: {mass}')
    doc.save(f'{shape}_report.docx')

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet['A1'] = 'Фигура'
    sheet['B1'] = 'Объем'
    sheet['C1'] = 'Площадь поверхности'
    sheet['D1'] = 'Масса'
    sheet.append([shape, volume, surface_area, mass])
    workbook.save(f'{shape}_report.xlsx')

    messagebox.showinfo("Сохранено", "Отчет сохранен в формате .docx и .xlsx.")

def update_fields(*args):
    if geometry_type.get() == "Параллелепипед":
        label_a.config(text="Длина (a):")
        label_b.config(text="Ширина (b):")
        label_h.config(text="Высота (h):")
        label_r.config(text="")
        entry_b.config(state='normal')
        entry_h.config(state='normal')
        entry_r.config(state='disabled')
    elif geometry_type.get() == "Тетраэдр":
        label_a.config(text="Длина ребра (a):")
        label_b.config(text="")
        label_h.config(text="")
        entry_b.config(state='disabled')
        entry_h.config(state='disabled')
        entry_r.config(state='disabled')
    elif geometry_type.get() == "Шар":
        label_a.config(text="")
        label_b.config(text="")
        label_h.config(text="")
        label_r.config(text="Радиус:")
        entry_a.config(state='disabled')
        entry_b.config(state='disabled')
        entry_h.config(state='disabled')
        entry_r.config(state='normal')

root = tk.Tk()
root.title("Геометрический калькулятор")

geometry_type = tk.StringVar(value="Параллелепипед")
geometry_type.trace_add('write', update_fields)

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

label_shape = tk.Label(frame, text="Выберите фигуру:")
label_shape.grid(row=0, column=0)

menu_shape = tk.OptionMenu(frame, geometry_type, "Параллелепипед", "Тетраэдр", "Шар")
menu_shape.grid(row=0, column=1)

label_a = tk.Label(frame, text="Длина (a):")
label_a.grid(row=1, column=0)

entry_a = tk.Entry(frame)
entry_a.grid(row=1, column=1)

label_b = tk.Label(frame, text="Ширина (b):")
label_b.grid(row=2, column=0)

entry_b = tk.Entry(frame)
entry_b.grid(row=2, column=1)

label_h = tk.Label(frame, text="Высота (h):")
label_h.grid(row=3, column=0)

entry_h = tk.Entry(frame)
entry_h.grid(row=3, column=1)

label_r = tk.Label(frame, text="Радиус (r):")
label_r.grid(row=4, column=0)

entry_r = tk.Entry(frame)
entry_r.grid(row=4, column=1)

label_density = tk.Label(frame, text="Плотность материала:")
label_density.grid(row=5, column=0)

entry_density = tk.Entry(frame)
entry_density.grid(row=5, column=1)

calculate_button = tk.Button(frame, text="Рассчитать", command=calculate)
calculate_button.grid(row=6, column=0, columnspan=2)

save_button = tk.Button(frame, text="Сохранить отчет", command=save_report)
save_button.grid(row=7, column=0, columnspan=2)

result_label = tk.Label(frame, text="")
result_label.grid(row=8, column=0, columnspan=2)


from pymongo import MongoClient
from datetime import datetime

# Функция для подключения к MongoDB
def get_mongo_client():
    return MongoClient("mongodb://localhost:27017/")

def save_to_mongodb():
    current_result = result_label.cget("text")
    if not current_result or "Объем:" not in current_result:
        messagebox.showwarning("Предупреждение", "Нет данных для сохранения. Сначала выполните расчет.")
        return
    
    try:
        shape = geometry_type.get()
        parts = current_result.split("\n")
        if len(parts) < 3:
            raise ValueError("Неверный формат результатов")
            
        volume = parts[0].split(":")[1].strip() if ":" in parts[0] else ""
        surface_area = parts[1].split(":")[1].strip() if ":" in parts[1] else ""
        mass = parts[2].split(":")[1].strip() if ":" in parts[2] else ""

        client = MongoClient("mongodb://localhost:27017/")
        db = client["geocalc"]
        collection = db["calculations"]
        
        document = {
            "shape": shape,
            "volume": volume,
            "surface_area": surface_area,
            "mass": mass,
            "calculation_time": datetime.now()
        }
        
        collection.insert_one(document)
        messagebox.showinfo("Успех", "Данные сохранены в MongoDB")
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении в MongoDB: {str(e)}")
    finally:
        if 'client' in locals():
            client.close()


mongo_button = tk.Button(frame, text="Сохранить в MongoDB", command=save_to_mongodb)
mongo_button.grid(row=10, column=0, columnspan=2)
# docker exec geo-calc-mongo mongosh geocalc --eval "db.calculations.find().pretty()"
update_fields()

root.mainloop()
