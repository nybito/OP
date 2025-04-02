import kivy
import mod
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.dropdown import DropDown
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from docx import Document
import openpyxl

# Импортируйте ваши классы для фигур здесь
# from mod import Para, Tetra, Shar

class GeometryCalculator(BoxLayout):
    def __init__(self, **kwargs):
        super(GeometryCalculator, self).__init__(**kwargs)
        self.orientation = 'vertical'

        self.shape_spinner = Spinner(
            text='Параллелепипед',
            values=('Параллелепипед', 'Тетраэдр', 'Шар'),
            size_hint=(None, None),
            size=(200, 44)
        )
        self.shape_spinner.bind(text=self.update_fields)

        self.add_widget(Label(text="Выберите фигуру:"))
        self.add_widget(self.shape_spinner)

        self.entry_a = TextInput(hint_text='Длина (a)', multiline=False)
        self.entry_b = TextInput(hint_text='Ширина (b)', multiline=False)
        self.entry_h = TextInput(hint_text='Высота (h)', multiline=False)
        self.entry_r = TextInput(hint_text='Радиус (r)', multiline=False)
        self.entry_density = TextInput(hint_text='Плотность материала', multiline=False)

        self.add_widget(self.entry_a)
        self.add_widget(self.entry_b)
        self.add_widget(self.entry_h)
        self.add_widget(self.entry_r)
        self.add_widget(self.entry_density)

        calculate_button = Button(text="Рассчитать")
        calculate_button.bind(on_press=self.calculate)
        
        save_button = Button(text="Сохранить отчет")
        save_button.bind(on_press=self.save_report)

        self.result_label = Label(text="")
        
        self.add_widget(calculate_button)
        self.add_widget(save_button)
        self.add_widget(self.result_label)

        # Инициализация полей в зависимости от выбранной фигуры
        self.update_fields()

    def update_fields(self, *args):
        shape = self.shape_spinner.text
        
        if shape == "Параллелепипед":
            self.entry_a.hint_text = "Длина (a)"
            self.entry_b.hint_text = "Ширина (b)"
            self.entry_h.hint_text = "Высота (h)"
            self.entry_r.disabled = True
            
            # Enable the relevant fields
            for entry in [self.entry_a, self.entry_b, self.entry_h]:
                entry.disabled = False
            
            for entry in [self.entry_r]:
                entry.text = ""
                entry.disabled = True

        elif shape == "Тетраэдр":
            self.entry_a.hint_text = "Длина ребра (a)"
            for entry in [self.entry_b, self.entry_h, self.entry_r]:
                entry.text = ""
                entry.disabled = True
            
            # Enable the relevant field
            for entry in [self.entry_a]:
                entry.disabled = False

        elif shape == "Шар":
            for entry in [self.entry_a, self.entry_b, self.entry_h]:
                entry.text = ""
                entry.disabled = True
            
            # Enable the relevant field
            for entry in [self.entry_r]:
                entry.hint_text = "Радиус (r)"
                entry.disabled = False

    def calculate(self, instance):
        shape = self.shape_spinner.text
        
        try:
            if shape == "Параллелепипед":
                a = float(self.entry_a.text)
                b = float(self.entry_b.text)
                h = float(self.entry_h.text)
                ro = float(self.entry_density.text)
                
                Fig = mod.Para(a,b,h,ro)  # Предполагается наличие класса Para в модуле mod.
                
            elif shape == "Тетраэдр":
                a = float(self.entry_a.text)
                ro = float(self.entry_density.text)
                
                Fig = mod.Tetra(a,ro)  # Предполагается наличие класса Tetra в модуле mod.

            elif shape == "Шар":
                r = float(self.entry_r.text)
                ro = float(self.entry_density.text)

                Fig = mod.Shar(r,ro)  # Предполагается наличие класса Shar в модуле mod.

            volume_str=f"Объем: {Fig.V}"
            surface_area_str=f"Площадь поверхности: {Fig.P}"
            mass_str=f"Масса: {Fig.M}"
            
            result_text=f"{volume_str}\n{surface_area_str}\n{mass_str}"
            
            # Обновление результата на экране
            self.result_label.text=result_text

        except ValueError as ve:
            popup=Popup(title='Ошибка', content=Label(text=str(ve)), size_hint=(None,None), size=(400,200))
            popup.open()

    def save_report(self, instance):
        
         # Получение значений из результата 
         result_lines=self.result_label.text.split("\n")
         volume=result_lines[0].split(":")[1].strip()
         surface_area=result_lines[1].split(":")[1].strip()
         mass=result_lines[2].split(":")[1].strip()
         
         shape=self.shape_spinner.text

         doc=Document()
         doc.add_heading(f'Отчет по {shape}', level=1)
         doc.add_paragraph(f'Объем: {volume}')
         doc.add_paragraph(f'Площадь поверхности: {surface_area}')
         doc.add_paragraph(f'Mасса: {mass}')
         doc.save(f'{shape}_report.docx')

         workbook=openpyxl.Workbook()
         sheet=workbook.active
         sheet['A1']='Фигура'
         sheet['B1']='Объем'
         sheet['C1']='Площадь поверхности'
         sheet['D1']='Масса'
         
         sheet.append([shape, volume, surface_area, mass])
         
         workbook.save(f'{shape}_report.xlsx')

         popup=Popup(title='Сохранено', content=Label(text="Отчет сохранен в формате .docx и .xlsx."), size_hint=(None,None), size=(400,200))
         popup.open()

class GeometryApp(App):
    def build(self):
        return GeometryCalculator()

if __name__ == '__main__':
    GeometryApp().run()
